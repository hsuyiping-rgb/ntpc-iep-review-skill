import argparse
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_NAME = "Microsoft JhengHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)


def set_font(run, size=11, bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text="", size=11, bold=False, color=None, before=0, after=6, style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    add_paragraph(
        doc,
        text,
        size=16 if level == 1 else 13,
        bold=True,
        color=BLUE if level == 1 else DARK_BLUE,
        before=16 if level == 1 else 12,
        after=8 if level == 1 else 6,
    )


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_font(run)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_font(run)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for style_name in ["Normal", "List Number", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("IEP Review")
    set_font(run, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("IEP 審查")
    set_font(run, size=9, color=GRAY)


def verify_docx(path):
    with zipfile.ZipFile(path) as package:
        names = set(package.namelist())
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise RuntimeError(f"{path} is not a valid DOCX package")

    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "???" in text:
        raise RuntimeError(f"{path} contains replacement question marks; check text encoding")
    return len([paragraph for paragraph in doc.paragraphs if paragraph.text.strip()])


def build(args):
    doc = Document()
    configure_document(doc)

    add_paragraph(doc, f"{args.student} IEP 審查意見", size=22, bold=True, color=DARK_BLUE, after=4)
    add_paragraph(doc, f"{args.year} 學年度 - {args.class_name}", size=13, color=GRAY, after=14)

    add_heading(doc, "學生資料")
    for line in [
        f"學生姓名：{args.student}",
        f"學年度：{args.year}",
        f"年級班級：{args.class_name}",
        f"審查結論：{args.result}",
    ]:
        add_paragraph(doc, line, after=4)

    add_heading(doc, "審查結論")
    add_paragraph(doc, f"建議：{args.result}", bold=True)
    if args.conclusion:
        add_paragraph(doc, args.conclusion)

    add_heading(doc, "主要補正意見")
    for item in args.corrections:
        add_number(doc, item)

    if args.keep:
        add_heading(doc, "可保留內容")
        for item in args.keep:
            add_bullet(doc, item)

    if args.sections:
        add_heading(doc, "已讀取頁面")
        for item in args.sections:
            add_bullet(doc, item)

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    paragraph_count = verify_docx(out)
    print(f"{out} ({paragraph_count} paragraphs)")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--student", required=True)
    parser.add_argument("--year", required=True)
    parser.add_argument("--class", dest="class_name", required=True)
    parser.add_argument("--result", required=True)
    parser.add_argument("--corrections", action="append", default=[])
    parser.add_argument("--keep", action="append", default=[])
    parser.add_argument("--sections", action="append", default=[])
    parser.add_argument("--conclusion", default="")
    parser.add_argument("--out", required=True)
    build(parser.parse_args())


if __name__ == "__main__":
    main()
